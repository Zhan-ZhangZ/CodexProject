"""DOCX parser for contracts, complaints, and other legal documents."""

from __future__ import annotations

import io

from docx import Document

from mcp_server.parsers.base import BaseParser, ParsedDocument, ensure_bytes, normalize_text


class DOCXParser(BaseParser):
    """Parse Word documents into paragraphs and table text."""

    supported_extensions = (".docx",)

    def parse(self, document_content: str | bytes, file_name: str) -> ParsedDocument:
        docx_bytes = ensure_bytes(document_content, file_name=file_name)
        document = Document(io.BytesIO(docx_bytes))

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        tables: list[str] = []

        for table_index, table in enumerate(document.tables, start=1):
            row_texts: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    row_texts.append(" | ".join(cells))
            if row_texts:
                tables.append(f"[表格{table_index}]\n" + "\n".join(row_texts))

        sections = []
        if paragraphs:
            sections.append("\n".join(paragraphs))
        if tables:
            sections.append("\n\n".join(tables))

        text = normalize_text("\n\n".join(sections))

        return ParsedDocument(
            file_name=file_name,
            file_type=".docx",
            text=text,
            metadata={
                "parser": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(document.tables),
                "character_count": len(text),
            },
        )
